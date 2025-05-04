import datetime as dt
import re
from io import BytesIO

from aiogram import Router, F
from aiogram.filters import Command, CommandStart
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from boto3.dynamodb.conditions import Key, Attr
from docx import Document

from data import cfg
from data.config_ydb import table_temp, table_bl, table_banned_user
from keyboards import keyboard
from lexicon import reminders

router = Router()
router.message.filter(F.chat.type.in_({"private"}))

lang_name = {'black_list_org': 'Название организации',
             'black_list_sotr': 'Ф.И.О. сотрудника'}


# Приветствие в личных сообщениях
@router.message(CommandStart() or Command(commands='help'))
async def send_welcome(message: Message):
    # This handler will be called when user sends `/start` or `/help` command
    await message.reply("Привет!\nЯ Робо-админ группы Геодезия_Работа и подработка!",
                        reply_markup=choice_keyboard(message.from_user.id))

@router.message(F.text == 'Правила группы')
async def rules(message: Message):
    rules_msg = 'Правила группы:\n\n'
    rules_msg = rules_msg + '\n\n'.join(reminders.reminders)

    await message.answer(f'{rules_msg}\n\n{reminders.reminders_full[0]}\n\n'
                         f'Рады приветствовать Вас!')


# ~~~~~~~~~~~~~~~Машина состояний~~~~~~~~~~~~~
# Начало по команде "Добавить"
@router.message(F.text == 'Добавить в черные списки')
async def bl_add_start(message: Message):
    await message.answer('Кого добавляем?',
                         reply_markup=keyboard.bl_add)


# Отмена процедуры добавления
@router.callback_query(F.data == 'cancel')
async def bl_add_cancel(call: CallbackQuery):
    try:
        table_temp.delete_item(
            Key={'id_user': call.from_user.id}
        )
        await call.message.edit_text('✅ Действие отменено.')
        await call.answer()
    except:
        print('Error')
        await call.message.edit_text('Ну и нахуй ты сюда тыкнул?')


# Выбор организации или сотрудника
@router.callback_query(F.data.startswith('bl_add_'))
async def bl_add_choice(call: CallbackQuery):

    if call.data == 'bl_add_org':
        call_choice = 'организацию'
        bl_base = 'black_list_org'
    else:
        call_choice = 'сотрудника'
        bl_base = 'black_list_sotr'
    table_temp.put_item(
        Item={
            'id_user': call.from_user.id,
            'body': {
                'state': 'bl_name',
                'bl_base': bl_base,
                'date': dt.datetime.now().strftime('%Y%m%d%H%M'),
                'name': '',
                'contact': '',
                'comment': '',
                'done': False
            }
        })
    await call.answer()
    await call.message.edit_text(f'ОК. \n\n'
                                 f'Добавляем в Ч.С. {call_choice}.\n'
                                 f'Нужно будет последовательно заполнить:\n\n'
                                 f'{lang_name[bl_base]}:\n>> ✏️\n'
                                 f'Контакты (контактное лицо):\n>> ✏️\n'
                                 f'Комментарий (причина добавления):\n>> ✏️\n\n'
                                 f'Изменить опечатки или дополнить информацию можно будет на последнем этапе '
                                 f'перед сохранением.\n'
                                 f'📝 Напиши {lang_name[bl_base]} или нажми "⛔️Отмена".',
                                 reply_markup=keyboard.bl_add_cancel)


# Запрос названия
def bl_add_name(message, state='bl_save'):
    table_temp.update_item(
        Key={
            'id_user': message.from_user.id
        },
        UpdateExpression="set body.state = :s, body.name = :n",
        ExpressionAttributeValues={
            ':s': state,
            ':n': message.text
        },
        ReturnValues="UPDATED_NEW"
    )
    return


# Запрос контактных данных
def bl_add_contact(message, state='bl_save'):
    table_temp.update_item(
        Key={
            'id_user': message.from_user.id
        },
        UpdateExpression="set body.state = :s, body.contact = :c",
        ExpressionAttributeValues={
            ':s': state,
            ':c': message.text
        },
        ReturnValues="UPDATED_NEW"
    )
    return


# Запрос комментария и сохранения.
def bl_add_comment(message):
    table_temp.update_item(
        Key={
            'id_user': message.from_user.id
        },
        UpdateExpression="set body.state = :s, body.comment = :c, body.done = :d",
        ExpressionAttributeValues={
            ':s': 'bl_save',
            ':c': message.text,
            ':d': True

        },
        ReturnValues="UPDATED_NEW"
    )
    return

# Проверка написанной записи и этап сохранения
@router.callback_query(F.data.startswith('bl_save_'))
async def bl_add_save(call: CallbackQuery):
    if call.data == 'bl_save_Save':
        response = table_temp.get_item(Key={'id_user': call.from_user.id})
        table_bl.put_item(
            Item={
                'from_user': response['Item']['id_user'],
                'date': response['Item']['body']['date'],
                'body': {
                    'bl_base': response['Item']['body']['bl_base'],
                    'name': response['Item']['body']['name'],
                    'contact': response['Item']['body']['contact'],
                    'comment': response['Item']['body']['comment']
                }
            }
        )
        await call.message.answer(text= f'З{call.message.text[7:-48]}\n\n'
                                                  f'✅ Успешно добавлена в Ч.С.',
                               reply_markup=choice_keyboard(call.from_user.id))
        try:
            await call.message.delete()
            table_temp.delete_item(
                Key={'id_user': call.from_user.id}
            )
        except:
            print("wtf?")
        await call.answer()
    else:
        text = call.message.text[:-47]
        await call.message.edit_text(f'{text}\n'
                                     f'✏️Что меняем?',
                                     reply_markup=keyboard.bl_add_edit)

# Режим изменения вносимых в Ч.С. данных
@router.callback_query(F.data.startswith('bl_edit_'))
async def bl_add_edit(call: CallbackQuery):
    choise_text = ''
    if call.data == 'bl_edit_name':
        table_temp.update_item(
            Key={
                'id_user': call.from_user.id
            },
            UpdateExpression="set body.state = :s",
            ExpressionAttributeValues={
                ':s': 'bl_name'
            },
            ReturnValues="UPDATED_NEW"
        )
        choise_text = 'новое Название или Ф.И.О.'
    elif call.data == 'bl_edit_contact':
        table_temp.update_item(
            Key={
                'id_user': call.from_user.id
            },
            UpdateExpression="set body.state = :s",
            ExpressionAttributeValues={
                ':s': 'bl_contact'
            },
            ReturnValues="UPDATED_NEW"
        )
        choise_text = 'новые данные о контактах.'
    elif call.data == 'bl_edit_comment':
        table_temp.update_item(
            Key={
                'id_user': call.from_user.id
            },
            UpdateExpression="set body.state = :s",
            ExpressionAttributeValues={
                ':s': 'bl_comment'
            },
            ReturnValues="UPDATED_NEW"
        )
        choise_text = 'новую причину добавления.'
    await call.message.edit_text(f'Хорошо, изменим. Жду от тебя {choise_text}')


# Начало по команде "Поиск по базе данных"
@router.message(F.text.startswith('Поиск'))
async def bl_search(message: Message):
    await message.answer('По какой базе будем искать?',
                         reply_markup=keyboard.bl_search)

# Выбор базы для поиска в Ч.С.
@router.callback_query(F.data.startswith('black_list_'))
async def bl_search_list(call: CallbackQuery):
    if call.data == 'black_list_my':
        response = table_bl.query(
            KeyConditionExpression=Key('from_user').eq(call.from_user.id)
        )
        if len(response['Items']) != 0:
            await call.message.edit_text(f'Твои записи:')
            for msg in response['Items']:
                await call.message.answer(f"----------------------------------\n"
                                          f"{lang_name[msg['body']['bl_base']]}:\n"
                                          f"{msg['body']['name']}\n\n"
                                          f"Контакты (контактное лицо):\n"
                                          f"{msg['body']['contact']}\n\n"
                                          f"Комментарий (причина добавления):\n"
                                          f"{msg['body']['comment']}\n\n"
                                          f"Дата добавления:\n"
                                          f"{str(dt.datetime.strptime(msg['date'], '%Y%m%d%H%M'))}\n\n",
                                          reply_markup=keyboard.bl_delete)

        else:
            await call.message.edit_text(f'Записей не найдено.')
    else:
        await call.answer()
        chioce = ''
        table_temp.put_item(
            Item={
                'id_user': call.from_user.id,
                'body': {
                    'state': 'search',
                    'search_list': call.data
                }
            })
        if call.data == 'black_list_org':
            chioce = 'организаций'
        elif call.data == 'black_list_sotr':
            chioce = 'сотрудников'
        await call.message.edit_text(f'Выполняется поиск по базе {chioce}.\n\n '
                                     f'Отправь кого ищем или нажми "⛔️Отмена"',
                                     reply_markup=keyboard.bl_full_base)

# Поиск по черному списку
async def black_list_scan(message, base):
    response_scan = table_bl.scan(
        FilterExpression=Attr('body.bl_base').eq(base)
    )
    result_scan = 0
    for item in response_scan['Items']:
        if item['body']['name'].lower().find(message.text.lower()) != -1:
            result_scan += 1
            key_del = ''
            markup = None
            if message.from_user.id in cfg.list_admin:
                key_del = f'*****\n' \
                          f'Код удаления записи:\n' \
                          f'DFBD*{item["from_user"]}*{item["date"]}\n\n'
                markup = keyboard.bl_delete_admin
            result_scan_item = f"{lang_name[item['body']['bl_base']]}:\n" \
                               f"{item['body']['name']}\n\n" \
                               f"Контакты (контактное лицо):\n" \
                               f"{item['body']['contact']}\n\n" \
                               f"Комментарий (причина добавления):\n" \
                               f"{item['body']['comment']}\n\n" \
                               f"Дата добавления:\n" \
                               f"{str(dt.datetime.strptime(item['date'], '%Y%m%d%H%M'))}\n\n{key_del}"
            await message.answer(result_scan_item, reply_markup=markup)

    if result_scan != 0:
        await message.answer(f'По запросу "{message.text}" найдено записей: {result_scan}',
                             reply_markup=choice_keyboard(message.from_user.id))
    else:
        await message.answer(f'По запросу "{message.text}" ничего не найдено.',
                             reply_markup=choice_keyboard(message.from_user.id))

    try:
        table_temp.delete_item(
            Key={'id_user': message.from_user.id}
        )

    except:
        print("wtf?")
    return result_scan

# Команда для удаления из Ч.С. пользователем
@router.callback_query(F.data.startswith('bl_delete_'))
async def bl_delete_user(call: CallbackQuery):
    await call.answer()
    if call.data == 'bl_delete_user':
        data = call.message.text.split('Дата добавления:')[1]
        try:
            table_bl.delete_item(
                Key={'from_user': call.from_user.id,
                     'date': re.sub(r'\W', '', data)[:-2]}
            )
            await call.message.edit_text('Запись удалена за базы')
        except:
            await call.message.edit_text('Что-то пошло не так.')
    elif call.data == 'bl_delete_admin' and call.from_user.id in cfg.list_admin:
        await call.message.edit_text(call.message.text, reply_markup=keyboard.dfbd_admin)

# Формирование и отправка ТХТ файла с Ч.С.
@router.callback_query(F.data == 'full_base')
async def file_send(call: CallbackQuery):
    await call.answer()
    base_scan = table_temp.query(KeyConditionExpression=Key('id_user').eq(call.from_user.id))
    if base_scan:
        scan = base_scan['Items'][0]['body']['search_list']
        response_scan = table_bl.scan(
            FilterExpression=Attr('body.bl_base').eq(scan)
        )
        choice = ''
        if scan == 'black_list_org':
            choice = 'организаций'
        elif scan == 'black_list_sotr':
            choice = 'сотрудников'

        document = Document()
        document.add_heading(f'Черный список {choice}:\n\n')
        if call.from_user.id in cfg.list_admin:
            document.add_paragraph(
                "Инструкция для Администраторов по удалению записей:\n"
                "Для удаления записи нужно скопировать 'Код удаления записи' "
                "начинающийся с DFDB* "
                f"и отправить данный код Боту_администратору ({cfg.url_bot})\n\n"
            )
        for item in response_scan['Items']:
            key_del = ''
            if call.from_user.id in cfg.list_admin:
                key_del = f'*****\n' \
                          f'Код удаления записи:\n' \
                          f'DFBD*{item["from_user"]}*{item["date"]}\n\n'
            # Добавляем раздел для каждой записи
            document.add_heading('----------------------------------', level=1)
            document.add_paragraph(f"{lang_name[item['body']['bl_base']]}:")
            document.add_paragraph(f"{item['body']['name']}")

            document.add_paragraph("Контакты (контактное лицо):")
            document.add_paragraph(f"{item['body']['contact']}")

            document.add_paragraph("Комментарий (причина добавления):")
            document.add_paragraph(f"{item['body']['comment']}")

            document.add_paragraph("Дата добавления:")
            document.add_paragraph(
                f"{str(dt.datetime.strptime(item['date'], '%Y%m%d%H%M'))}"
            )
            document.add_paragraph(key_del)


        file = BytesIO()
        file.name = f'full_base_{choice}.docx'
        document.save(file)
        file.seek(0)

        await call.message.edit_text(f'Черный список {choice}:')
        document_file = BufferedInputFile(file.getvalue(), filename=file.name)
        await call.message.answer_document(document=document_file)

        try:
            table_temp.delete_item(
                Key={'id_user': call.from_user.id}
            )
        except:
            print("Ошибка при удалении")
    else:
        await call.message.edit_text('Что-то пошло не так.\nПопробуйте начать сначала.')


@router.message(F.text == 'Узнать причину бана')
async def baned_user(message: Message):
    response = table_banned_user.query(
        KeyConditionExpression=Key('id_user').eq(message.from_user.id)
    )
    if len(response['Items']) != 0:

        body = response['Items'][0]['body']

        await message.answer(f'Сообщение отправленное '
                             f'{body["date"][:-6]} UTC:\n'
                             f'{body["message_text"]}\n\n'
                             f'Комментарий:\n'
                             f'{body["comment"]}\n\n'
                             f'*Перечитайте правила группы.\n'
                             f'**Если не видите причины бана в правилах, перечитайте еще раз.\n'
                             f'Если не нашли подходящего, напишите админам, попробуем разобраться (но это не точно).')
    else:
        await message.answer('Я не знаю. Возможно вы попали в бан еще до создания этого мира.')


# Функционал в ЛС
@router.message()
async def private_dialog(message: Message):
    response = table_temp.query(KeyConditionExpression=Key('id_user').eq(message.from_user.id))
    if len(response['Items']) != 0:
        state = response['Items'][0]['body']['state']
        if state == 'search':
            await black_list_scan(message, response['Items'][0]['body']['search_list'])

        else:
            done = response['Items'][0]['body']['done']
            name = response['Items'][0]['body']['name']
            contact = response['Items'][0]['body']['contact']
            comment = response['Items'][0]['body']['comment']
            step_message = 'Проверяем. Если все верно, нажми "💾 Сохранить".'
            step_button = keyboard.bl_add_save
            if state == 'bl_name':
                if not done:
                    bl_add_name(message, 'bl_contact')
                    step_message = '📝 Напиши контактные данные или нажми "⛔️Отмена".'
                    step_button = keyboard.bl_add_cancel
                else:
                    bl_add_name(message)
                name = message.text
            elif state == 'bl_contact':
                if not done:
                    bl_add_contact(message, 'bl_comment')
                    step_message = '📝 Напиши причину добавления. Чем подробнее, тем лучше. Ну или нажми "⛔️Отмена".'
                    step_button = keyboard.bl_add_cancel
                else:
                    bl_add_contact(message)
                contact = message.text
            elif state == 'bl_comment':
                bl_add_comment(message)
                comment = message.text
            await message.answer(f'Новая запись:\n\n'
                                 f"{lang_name[response['Items'][0]['body']['bl_base']]}:\n"
                                 f">> {name}\n"
                                 f"Контакты (контактное лицо):\n"
                                 f">> {contact}\n"
                                 f"Комментарий (причина добавления):\n"
                                 f">> {comment}\n\n"
                                 f"{step_message}", reply_markup=step_button)
    else:
        await message.answer(f'Привет. Я тебя не понял. Нажми на кнопку ниже.',
                             reply_markup=choice_keyboard(message.from_user.id))



# Клавиатура админам
def choice_keyboard(id_user):
    if id_user in cfg.list_admin:
        keyboard_answer = keyboard.keyboard_private_admin
    else:
        keyboard_answer = keyboard.keyboard_private
    return keyboard_answer
